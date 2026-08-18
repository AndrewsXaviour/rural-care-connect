from markdown import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

def md_to_docx(md_file, docx_file):
    """Convert a Markdown file to DOCX using python-docx."""
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Process line by line for better control
    lines = md_content.split('\n')
    in_code_block = False
    code_lines = []
    in_table = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 30, 30)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Skip empty lines
        if not stripped:
            continue
        
        # Headings
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            text = stripped.lstrip('#').strip()
            heading = doc.add_heading(text, level=min(level, 4))
            continue
        
        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', stripped):
            doc.add_paragraph('_' * 60).paragraph_format.space_after = Pt(6)
            continue
        
        # Unordered list items
        if re.match(r'^[\s]*[-*+]\s+', stripped):
            text = re.sub(r'^[\s]*[-*+]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_inline_formatting(p, text)
            continue
        
        # Ordered list items
        if re.match(r'^\s*\d+[.)]\s+', stripped):
            text = re.sub(r'^\s*\d+[.)]\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_inline_formatting(p, text)
            continue
        
        # Bold text (common patterns like **text**)
        # Regular paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_inline_formatting(p, stripped)
    
    doc.save(docx_file)
    print(f"✓ Converted: {md_file} -> {docx_file}")

def _add_inline_formatting(paragraph, text):
    """Add text with bold/italic formatting support."""
    # Split on **bold** or *italic* patterns
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold and italic
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        # Bold
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Italic
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        # Inline code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        # Links [text](url)
        elif part.startswith('[') and '](' in part:
            link_text = part[1:part.index('](')]
            link_url = part[part.index('](')+2:part.index(')')]
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
        else:
            paragraph.add_run(part)


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    files = [
        ('LEARNING_ROADMAP.md', 'LEARNING_ROADMAP.docx'),
        ('RESUME_CONTENT.md', 'RESUME_CONTENT.docx'),
    ]
    
    for md_name, docx_name in files:
        md_path = os.path.join(base_dir, md_name)
        docx_path = os.path.join(base_dir, docx_name)
        if os.path.exists(md_path):
            md_to_docx(md_path, docx_path)
        else:
            print(f"⚠ File not found: {md_path}")
    
    print("\nDone! Both files converted.")